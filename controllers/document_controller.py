#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Contrôleur de gestion des documents pour l'application Vynal Docs Automator
"""

import os
import re
import shutil
import logging
import tkinter as tk
import customtkinter as ctk
from tkinter import messagebox, filedialog, simpledialog
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_analyzer import DocumentAnalyzer

logger = logging.getLogger("VynalDocsAutomator.DocumentController")

class DialogUtils:
    """
    Utilitaires pour créer des boîtes de dialogue cohérentes dans l'application
    """
    
    @staticmethod
    def show_confirmation(parent, title, message, on_yes=None, on_no=None):
        """
        Affiche une boîte de dialogue de confirmation
        
        Args:
            parent: Widget parent
            title: Titre de la boîte de dialogue
            message: Message à afficher
            on_yes: Fonction à appeler si l'utilisateur confirme
            on_no: Fonction à appeler si l'utilisateur annule
            
        Returns:
            bool: True si confirmé, False sinon
        """
        dialog = ctk.CTkToplevel(parent)
        dialog.title(title)
        dialog.geometry("400x200")
        dialog.resizable(False, False)
        dialog.grab_set()
        dialog.focus_set()
        
        # Centrer la fenêtre
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
        y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # Résultat par défaut
        result = [False]
        
        # Cadre principal
        frame = ctk.CTkFrame(dialog, fg_color="transparent")
        frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
        
        # Titre avec icône
        ctk.CTkLabel(
            frame,
            text=f"⚠️ {title}",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=(0, 10))
        
        # Message
        ctk.CTkLabel(
            frame,
            text=message,
            wraplength=360
        ).pack(pady=10)
        
        # Cadre pour les boutons
        button_frame = ctk.CTkFrame(frame, fg_color="transparent")
        button_frame.pack(pady=10)
        
        # Fonctions de callback
        def yes_action():
            result[0] = True
            dialog.destroy()
            if on_yes:
                on_yes()
        
        def no_action():
            result[0] = False
            dialog.destroy()
            if on_no:
                on_no()
        
        # Bouton Non
        ctk.CTkButton(
            button_frame,
            text="Non",
            command=no_action,
            width=100,
            fg_color="#e74c3c",
            hover_color="#c0392b"
        ).pack(side=ctk.LEFT, padx=10)
        
        # Bouton Oui
        ctk.CTkButton(
            button_frame,
            text="Oui",
            command=yes_action,
            width=100,
            fg_color="#2ecc71",
            hover_color="#27ae60"
        ).pack(side=ctk.LEFT, padx=10)
        
        # Attendre que la fenêtre soit fermée
        parent.wait_window(dialog)
        
        return result[0]
    
    @staticmethod
    def show_message(parent, title, message, message_type="info"):
        """
        Affiche une boîte de dialogue avec un message
        
        Args:
            parent: Widget parent
            title: Titre de la boîte de dialogue
            message: Message à afficher
            message_type: Type de message ('info', 'error', 'warning', 'success')
        """
        dialog = ctk.CTkToplevel(parent)
        dialog.title(title)
        dialog.geometry("400x200")
        dialog.resizable(False, False)
        dialog.grab_set()
        dialog.focus_set()
        
        # Centrer la fenêtre
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
        y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # Icône selon le type
        icon = "ℹ️"
        button_color = "#3498db"
        hover_color = "#2980b9"
        
        if message_type == "error":
            icon = "❌"
            button_color = "#e74c3c"
            hover_color = "#c0392b"
        elif message_type == "warning":
            icon = "⚠️"
            button_color = "#f39c12"
            hover_color = "#d35400"
        elif message_type == "success":
            icon = "✅"
            button_color = "#2ecc71"
            hover_color = "#27ae60"
        
        # Cadre principal
        frame = ctk.CTkFrame(dialog, fg_color="transparent")
        frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
        
        # Titre avec icône
        ctk.CTkLabel(
            frame,
            text=f"{icon} {title}",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=(0, 10))
        
        # Message
        ctk.CTkLabel(
            frame,
            text=message,
            wraplength=360
        ).pack(pady=10)
        
        # Bouton OK
        ctk.CTkButton(
            frame,
            text="OK",
            command=dialog.destroy,
            width=100,
            fg_color=button_color,
            hover_color=hover_color
        ).pack(pady=10)

class DocumentController:
    """
    Contrôleur de gestion des documents
    Gère la logique métier liée aux documents
    """
    
    def __init__(self, app_model, document_view):
        """
        Initialise le contrôleur des documents
        
        Args:
            app_model: Modèle de l'application
            document_view: Vue de gestion des documents
        """
        self.model = app_model
        self.view = document_view
        
        # Initialiser l'analyseur de documents
        self.document_analyzer = DocumentAnalyzer()
        
        # Connecter les événements de la vue aux méthodes du contrôleur
        self.connect_events()
        
        logger.info("DocumentController initialisé")
    
    def connect_events(self):
        """
        Connecte les événements de la vue aux méthodes du contrôleur
        """
        # Remplacer les méthodes de la vue par les méthodes du contrôleur
        self.view.new_document = self.new_document
        self.view.open_document = self.open_document
        self.view.download_document = self.download_document
        self.view.filter_documents = self.filter_documents
        
        logger.info("Événements de DocumentView connectés")
    
    def handle_auto_fill(self, template_id: str, client_id: str = None) -> bool:
        """
        Gère l'auto-remplissage d'un document
        
        Args:
            template_id: ID du modèle à utiliser
            client_id: ID du client (optionnel)
            
        Returns:
            bool: True si l'auto-remplissage a réussi, False sinon
        """
        try:
            # Récupérer le modèle
            template = self.model.get_template(template_id)
            if not template:
                DialogUtils.show_message(self.view.parent, "Erreur", "Modèle non trouvé", "error")
                return False
            
            # Récupérer le client
            client = None
            if client_id:
                client = self.model.get_client(client_id)
            else:
                # Si aucun client n'est spécifié, utiliser le client sélectionné
                selected_client = self.view.get_selected_client()
                if selected_client:
                    client = self.model.get_client(selected_client)
            
            if not client:
                DialogUtils.show_message(self.view.parent, "Attention", "Veuillez sélectionner un client", "warning")
                return False
            
            # Analyser le modèle
            template_path = os.path.join(self.model.paths['templates'], template['file_path'])
            if not os.path.exists(template_path):
                DialogUtils.show_message(self.view.parent, "Erreur", "Fichier modèle non trouvé", "error")
                return False
            
            # Analyser le document
            analysis_result = self.document_analyzer.analyze_document(template_path)
            
            if 'error' in analysis_result:
                DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de l'analyse: {analysis_result['error']}", "error")
                return False
            
            # Extraire les données pertinentes
            extracted_data = analysis_result.get('data', {})
            
            # Créer une boîte de dialogue pour l'auto-remplissage
            dialog = ctk.CTkToplevel(self.view.parent)
            dialog.title("Auto-remplissage")
            dialog.geometry("600x400")
            dialog.resizable(True, True)
            dialog.grab_set()
            dialog.focus_set()
            
            # Centrer la fenêtre
            dialog.update_idletasks()
            x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
            y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
            dialog.geometry(f"+{x}+{y}")
            
            # Cadre principal
            main_frame = ctk.CTkFrame(dialog)
            main_frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
            
            # Titre
            ctk.CTkLabel(
                main_frame,
                text="Souhaitez-vous automatiser le remplissage de ce document?",
                font=ctk.CTkFont(size=16, weight="bold")
            ).pack(pady=10)
            
            # Liste des champs extraits
            fields_frame = ctk.CTkFrame(main_frame)
            fields_frame.pack(fill=ctk.BOTH, expand=True, pady=10)
            
            # Créer une liste scrollable pour les champs
            canvas = tk.Canvas(fields_frame)
            scrollbar = ctk.CTkScrollbar(fields_frame, orientation="vertical", command=canvas.yview)
            scrollable_frame = ctk.CTkFrame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            # Afficher les champs extraits
            for section, fields in extracted_data.items():
                if isinstance(fields, dict):
                    ctk.CTkLabel(
                        scrollable_frame,
                        text=section.title(),
                        font=ctk.CTkFont(weight="bold")
                    ).pack(pady=(10, 5), padx=10, anchor="w")
                    
                    for field, value in fields.items():
                        if value:
                            ctk.CTkLabel(
                                scrollable_frame,
                                text=f"{field}: {value}",
                                wraplength=400
                            ).pack(pady=2, padx=20, anchor="w")
            
            # Pack le canvas et la scrollbar
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Boutons
            button_frame = ctk.CTkFrame(main_frame)
            button_frame.pack(fill=ctk.X, pady=10)
            
            def on_confirm():
                # Créer le document avec les données extraites
                document_data = {
                    'template_id': template_id,
                    'client_id': client['id'],
                    'title': template['title'],
                    'type': template['type'],
                    'created_at': datetime.now().isoformat(),
                    'extracted_data': extracted_data
                }
                
                # Générer le document
                if self.generate_document_content(template, client, extracted_data):
                    dialog.destroy()
                    DialogUtils.show_message(self.view.parent, "Succès", "Document généré avec succès", "success")
                else:
                    dialog.destroy()
                    DialogUtils.show_message(self.view.parent, "Erreur", "Erreur lors de la génération du document", "error")
            
            def on_cancel():
                dialog.destroy()
            
            ctk.CTkButton(
                button_frame,
                text="Annuler",
                command=on_cancel,
                width=100,
                fg_color="#e74c3c",
                hover_color="#c0392b"
            ).pack(side=ctk.RIGHT, padx=10)
            
            ctk.CTkButton(
                button_frame,
                text="Confirmer",
                command=on_confirm,
                width=100,
                fg_color="#2ecc71",
                hover_color="#27ae60"
            ).pack(side=ctk.RIGHT, padx=10)
            
            # Attendre que la fenêtre soit fermée
            self.view.parent.wait_window(dialog)
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de l'auto-remplissage: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Une erreur est survenue: {str(e)}", "error")
            return False
    
    def new_document(self, template_id=None):
        """
        Crée un nouveau document à partir d'un modèle
        
        Args:
            template_id: ID du modèle à utiliser (optionnel)
        """
        # Vérifier s'il y a des modèles disponibles
        if not self.model.templates:
            DialogUtils.show_message(self.view.parent, "Attention", "Aucun modèle disponible. Veuillez d'abord créer un modèle.", "warning")
            return
        
        # Vérifier s'il y a des clients disponibles
        if not self.model.clients:
            DialogUtils.show_message(self.view.parent, "Attention", "Aucun client disponible. Veuillez d'abord ajouter un client.", "warning")
            return
        
        try:
            # Créer une nouvelle instance du formulaire de document
            from views.document_view import DocumentFormView
            form = DocumentFormView(self.view.parent, self.model, on_save_callback=self.view.update_view)
            
            # Si un modèle est spécifié, le charger
            if template_id:
                # Trouver le modèle
                template = next((t for t in self.model.templates if t.get("id") == template_id), None)
                if template:
                    # Pré-remplir le formulaire avec les données du modèle
                    form.template_data = template
                else:
                    DialogUtils.show_message(self.view.parent, "Erreur", "Modèle non trouvé", "error")
                    return
            
            logger.info(f"Nouveau document créé" + (f" avec le modèle {template_id}" if template_id else ""))
            
        except Exception as e:
            logger.error(f"Erreur lors de la création du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Une erreur est survenue: {str(e)}", "error")
    
    def open_document(self, document_id):
        """
        Ouvre un document pour le visualiser
        
        Args:
            document_id: ID du document à ouvrir
        """
        # Récupérer le document
        document = next((d for d in self.model.documents if d.get('id') == document_id), None)
        
        if not document:
            DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
            return
        
        # Vérifier si le fichier existe
        file_path = document.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            DialogUtils.show_message(self.view.parent, "Erreur", "Le fichier du document est introuvable", "error")
            return
        
        # Ouvrir le fichier avec l'application par défaut du système
        try:
            import subprocess
            
            if os.name == 'nt':  # Windows
                os.startfile(file_path)
            elif os.name == 'posix':  # macOS ou Linux
                subprocess.call(('open' if os.uname().sysname == 'Darwin' else 'xdg-open', file_path))
            
            logger.info(f"Document ouvert: {document_id} - {file_path}")
        except Exception as e:
            logger.error(f"Erreur lors de l'ouverture du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de l'ouverture du document: {str(e)}", "error")
    
    def download_document(self, document_id):
        """
        Télécharge (copie) un document vers un emplacement choisi par l'utilisateur
        
        Args:
            document_id: ID du document à télécharger
        """
        # Récupérer le document
        document = next((d for d in self.model.documents if d.get('id') == document_id), None)
        
        if not document:
            DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
            return
        
        # Vérifier si le fichier existe
        file_path = document.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            DialogUtils.show_message(self.view.parent, "Erreur", "Le fichier du document est introuvable", "error")
            return
        
        # Déterminer l'extension du fichier
        _, ext = os.path.splitext(file_path)
        
        # Ouvrir une boîte de dialogue pour choisir l'emplacement de sauvegarde
        dest_path = filedialog.asksaveasfilename(
            title="Enregistrer le document",
            defaultextension=ext,
            initialfile=os.path.basename(file_path),
            filetypes=[(f"Fichiers {ext.upper()}", f"*{ext}"), ("Tous les fichiers", "*.*")]
        )
        
        if not dest_path:
            return
        
        try:
            # Copier le fichier
            shutil.copy2(file_path, dest_path)
            
            logger.info(f"Document téléchargé: {document_id} - {dest_path}")
            DialogUtils.show_message(self.view.parent, "Succès", "Document téléchargé avec succès", "success")
        except Exception as e:
            logger.error(f"Erreur lors du téléchargement du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors du téléchargement du document: {str(e)}", "error")
    
    def filter_documents(self, *args):
        """
        Filtre les documents selon les critères sélectionnés
        """
        # Cette méthode est déjà implémentée dans la vue
        # Nous nous assurons juste qu'elle est appelée au bon moment
        self.view.update_view()
        logger.debug("Documents filtrés")
    
    def generate_document_content(self, template, client, variables):
        """
        Génère le contenu d'un document en remplaçant les variables
        
        Args:
            template: Modèle de document
            client: Informations du client
            variables: Valeurs des variables spécifiques
            
        Returns:
            str: Contenu du document avec les variables remplacées
        """
        # Récupérer le contenu du modèle
        content = template.get('content', "")
        
        # Remplacer les variables client
        content = content.replace("{client_name}", client.get('name', ''))
        content = content.replace("{client_company}", client.get('company', ''))
        content = content.replace("{client_email}", client.get('email', ''))
        content = content.replace("{client_phone}", client.get('phone', ''))
        content = content.replace("{client_address}", client.get('address', ''))
        
        # Remplacer les variables de l'entreprise
        content = content.replace("{company_name}", self.model.config.get("app.company_name", ""))
        
        # Remplacer la date
        content = content.replace("{date}", datetime.now().strftime("%Y-%m-%d"))
        
        # Remplacer les variables spécifiques
        for var_name, var_value in variables.items():
            content = content.replace(f"{{{var_name}}}", var_value)
        
        # Calculer automatiquement les montants TTC et TVA si nécessaire
        if "{montant_ttc}" in content and "montant" in variables:
            try:
                montant_ht = float(variables.get("montant", "0").replace(",", "."))
                montant_tva = montant_ht * 0.2  # TVA à 20%
                montant_ttc = montant_ht + montant_tva
                
                content = content.replace("{montant_tva}", f"{montant_tva:.2f}")
                content = content.replace("{montant_ttc}", f"{montant_ttc:.2f}")
            except ValueError:
                pass
        
        return content
    
    def generate_filename(self, doc_type, client_name, date):
        """
        Génère un nom de fichier pour le document
        
        Args:
            doc_type: Type de document
            client_name: Nom du client
            date: Date du document
            
        Returns:
            str: Nom de fichier normalisé
        """
        # Nettoyer les noms
        doc_type = self.clean_filename(doc_type)
        client_name = self.clean_filename(client_name)
        
        # Formatter la date
        date_format = self.model.config.get("document.date_format", "%Y-%m-%d")
        try:
            date_obj = datetime.strptime(date, "%Y-%m-%d")
            formatted_date = date_obj.strftime(date_format)
        except:
            formatted_date = datetime.now().strftime(date_format)
        
        # Construire le nom du fichier selon le pattern configuré
        pattern = self.model.config.get("document.filename_pattern", "{document_type}_{client_name}_{date}")
        
        filename = pattern.replace("{document_type}", doc_type)
        filename = filename.replace("{client_name}", client_name)
        filename = filename.replace("{date}", formatted_date)
        
        return filename
    
    def clean_filename(self, name):
        """
        Nettoie un nom pour qu'il soit utilisable dans un nom de fichier
        
        Args:
            name: Nom à nettoyer
            
        Returns:
            str: Nom nettoyé
        """
        # Supprimer les caractères spéciaux et remplacer les espaces par des underscores
        name = re.sub(r'[\\/*?:"<>|]', '', name)
        name = name.replace(' ', '_')
        name = name.replace('/', '_')
        name = name.replace('\\', '_')
        
        return name
    
    def generate_pdf(self, file_path, content, title, client, company_name):
        """
        Génère un fichier PDF
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document
            title: Titre du document
            client: Informations du client
            company_name: Nom de l'entreprise
        """
        try:
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Créer un PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4
            
            # En-tête
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, height - 50, company_name)
            
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 70, "Document généré par Vynal Docs Automator")
            c.drawString(50, height - 85, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            # Ligne séparatrice
            c.line(50, height - 95, width - 50, height - 95)
            
            # Titre
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 130, title)
            
            # Client
            c.setFont("Helvetica", 12)
            c.drawString(50, height - 160, f"Client: {client.get('name', '')}")
            if client.get('company'):
                c.drawString(50, height - 175, f"Entreprise: {client.get('company')}")
                y_offset = 190
            else:
                y_offset = 175
            
            c.drawString(50, height - y_offset, f"Email: {client.get('email', '')}")
            if client.get('phone'):
                y_offset += 15
                c.drawString(50, height - y_offset, f"Téléphone: {client.get('phone')}")
            
            # Ligne séparatrice
            y_offset += 15
            c.line(50, height - y_offset, width - 50, height - y_offset)
            
            # Contenu du document
            c.setFont("Helvetica", 10)
            y_position = height - y_offset - 25
            
            # Diviser le contenu en lignes
            lines = content.split('\n')
            
            for line in lines:
                if y_position < 50:
                    # Créer une nouvelle page
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y_position = height - 50
                
                # Écrire la ligne
                c.drawString(50, y_position, line)
                y_position -= 15
            
            # Finaliser le PDF
            c.showPage()
            c.save()
            
            logger.info(f"PDF généré: {file_path}")
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du PDF: {e}")
            raise
    
    def generate_docx(self, file_path, content, title, client, company_name):
        """
        Génère un fichier DOCX
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document
            title: Titre du document
            client: Informations du client
            company_name: Nom de l'entreprise
        """
        try:
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Créer un document Word
            doc = docx.Document()
            
            # En-tête
            header = doc.add_paragraph()
            header_run = header.add_run(company_name)
            header_run.bold = True
            header_run.font.size = Pt(14)
            
            doc.add_paragraph("Document généré par Vynal Docs Automator")
            doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            # Ligne séparatrice
            doc.add_paragraph("_______________________________________________________________")
            
            # Titre
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            
            # Client
            doc.add_paragraph(f"Client: {client.get('name', '')}")
            if client.get('company'):
                doc.add_paragraph(f"Entreprise: {client.get('company')}")
            doc.add_paragraph(f"Email: {client.get('email', '')}")
            if client.get('phone'):
                doc.add_paragraph(f"Téléphone: {client.get('phone')}")
            
            # Ligne séparatrice
            doc.add_paragraph("_______________________________________________________________")
            
            # Contenu
            # Diviser le contenu en paragraphes
            paragraphs = content.split('\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Pied de page
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_run = footer_para.add_run(f"© {datetime.now().year} {company_name} - Confidentiel")
            footer_run.font.size = Pt(8)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Enregistrer le document
            doc.save(file_path)
            
            logger.info(f"DOCX généré: {file_path}")
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du DOCX: {e}")
            raise
    
    def delete_document(self, document_id):
        """
        Supprime un document
        
        Args:
            document_id: ID du document à supprimer
        """
        # Récupérer le document
        document = next((d for d in self.model.documents if d.get('id') == document_id), None)
        
        if not document:
            DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
            return
        
        # Confirmer la suppression
        def delete_action():
            try:
                # Supprimer le fichier s'il existe
                file_path = document.get('file_path')
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)
                
                # Supprimer le document de la liste
                self.model.documents = [d for d in self.model.documents if d.get('id') != document_id]
                
                # Sauvegarder les modifications
                self.model.save_documents()
                
                # Ajouter l'activité
                self.model.add_activity("document", f"Document supprimé: {document.get('title')}")
                
                # Mettre à jour la vue
                self.view.update_view()
                
                logger.info(f"Document supprimé: {document_id} - {document.get('title')}")
                
                # Afficher un message de succès
                DialogUtils.show_message(self.view.parent, "Succès", "Document supprimé avec succès", "success")
            except Exception as e:
                logger.error(f"Erreur lors de la suppression du document: {e}")
                DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de la suppression du document: {str(e)}", "error")
        
        DialogUtils.show_confirmation(
            self.view.parent,
            "Confirmer la suppression",
            f"Êtes-vous sûr de vouloir supprimer le document '{document.get('title')}'?\n\nCette action est irréversible.",
            on_yes=delete_action
        )
    
    def export_document(self, document_id, format_type=None):
        """
        Exporte un document dans un autre format
        
        Args:
            document_id: ID du document à exporter
            format_type: Format d'exportation (pdf ou docx)
        """
        # Récupérer le document
        document = next((d for d in self.model.documents if d.get('id') == document_id), None)
        
        if not document:
            DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
            return
        
        # Vérifier si le fichier existe
        file_path = document.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            DialogUtils.show_message(self.view.parent, "Erreur", "Le fichier du document est introuvable", "error")
            return
        
        # Déterminer le format actuel
        _, current_ext = os.path.splitext(file_path)
        current_ext = current_ext.lower().strip('.')
        
        # Si aucun format spécifié, demander à l'utilisateur
        if not format_type:
            if current_ext == 'pdf':
                format_type = 'docx'
            else:
                format_type = 'pdf'
        
        # Vérifier que le format est différent du format actuel
        if format_type == current_ext:
            DialogUtils.show_message(self.view.parent, "Information", f"Le document est déjà au format {format_type.upper()}", "info")
            return
        
        # Déterminer le chemin du nouveau fichier
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        new_file_path = os.path.join(os.path.dirname(file_path), f"{base_name}.{format_type}")
        
        try:
            # Récupérer les informations nécessaires
            title = document.get('title', '')
            
            # Récupérer le client
            client_id = document.get('client_id')
            client = next((c for c in self.model.clients if c.get('id') == client_id), {})
            
            # Récupérer le modèle
            template_id = document.get('template_id')
            template = next((t for t in self.model.templates if t.get('id') == template_id), {})
            
            # Récupérer le contenu
            content = self.generate_document_content(
                template, 
                client, 
                document.get('variables', {})
            )
            
            # Générer le fichier selon le format
            company_name = self.model.config.get("app.company_name", "")
            
            if format_type == "pdf":
                self.generate_pdf(new_file_path, content, title, client, company_name)
            else:
                self.generate_docx(new_file_path, content, title, client, company_name)
            
            logger.info(f"Document exporté: {document_id} - {new_file_path}")
            
            # Demander à l'utilisateur s'il veut ouvrir le fichier
            def open_exported_file():
                # Ouvrir le fichier avec l'application par défaut du système
                if os.name == 'nt':  # Windows
                    os.startfile(new_file_path)
                elif os.name == 'posix':  # macOS ou Linux
                    import subprocess
                    subprocess.call(('open' if os.uname().sysname == 'Darwin' else 'xdg-open', new_file_path))
            
            DialogUtils.show_confirmation(
                self.view.parent,
                "Succès",
                f"Document exporté avec succès en format {format_type.upper()}.\n\nVoulez-vous l'ouvrir maintenant?",
                on_yes=open_exported_file
            )
            
        except Exception as e:
            logger.error(f"Erreur lors de l'exportation du document: {e}")
            DialogUtils.show_message(self.view.parent, "Erreur", f"Erreur lors de l'exportation du document: {str(e)}", "error")
    
    def update_document(self, document_id):
        """
        Met à jour un document existant
        
        Args:
            document_id: ID du document à mettre à jour
        """
        # Récupérer le document
        document = next((d for d in self.model.documents if d.get('id') == document_id), None)
        
        if not document:
            DialogUtils.show_message(self.view.parent, "Erreur", "Document non trouvé", "error")
            return
        
        # Récupérer le client et le modèle associés
        client_id = document.get('client_id')
        client = next((c for c in self.model.clients if c.get('id') == client_id), None)
        
        template_id = document.get('template_id')
        template = next((t for t in self.model.templates if t.get('id') == template_id), None)
        
        if not client or not template:
            DialogUtils.show_message(self.view.parent, "Erreur", "Client ou modèle introuvable", "error")
            return
        
        # Ouvrir une boîte de dialogue pour modifier les informations
        dialog = ctk.CTkToplevel(self.view.parent)
        dialog.title("Mettre à jour le document")
        dialog.geometry("550x600")
        dialog.resizable(False, False)
        dialog.grab_set()  # Modal
        dialog.focus_set()
        
        # Centrer la fenêtre
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
        y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # Créer un cadre principal avec défilement
        main_frame = ctk.CTkScrollableFrame(dialog)
        main_frame.pack(fill=ctk.BOTH, expand=True, padx=20, pady=20)
        
        # Informations du document
        info_frame = ctk.CTkFrame(main_frame)
        info_frame.pack(fill=ctk.X, pady=10)
        
        ctk.CTkLabel(info_frame, text="Informations du document", font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        info_list_frame = ctk.CTkFrame(info_frame)
        info_list_frame.pack(fill=ctk.X, padx=10, pady=10)
        
        # Titre
        title_row = ctk.CTkFrame(info_list_frame, fg_color="transparent")
        title_row.pack(fill=ctk.X, pady=5)
        ctk.CTkLabel(title_row, text="Titre:").pack(side=ctk.LEFT, padx=5)
        title_var = ctk.StringVar(value=document.get('title', ''))
        title_entry = ctk.CTkEntry(title_row, textvariable=title_var, width=300)
        title_entry.pack(side=ctk.LEFT, fill=ctk.X, expand=True, padx=5)
        
        # Date
        date_row = ctk.CTkFrame(info_list_frame, fg_color="transparent")
        date_row.pack(fill=ctk.X, pady=5)
        ctk.CTkLabel(date_row, text="Date:").pack(side=ctk.LEFT, padx=5)
        date_var = ctk.StringVar(value=document.get('date', ''))
        date_entry = ctk.CTkEntry(date_row, textvariable=date_var, width=200)
        date_entry.pack(side=ctk.LEFT, padx=5)
        
        # Description
        desc_row = ctk.CTkFrame(info_list_frame, fg_color="transparent")
        desc_row.pack(fill=ctk.X, pady=5)
        ctk.CTkLabel(desc_row, text="Description:").pack(side=ctk.LEFT, padx=5)
        description_frame = ctk.CTkFrame(desc_row, fg_color="transparent")
        description_frame.pack(side=ctk.LEFT, fill=ctk.X, expand=True, padx=5)
        description_text = ctk.CTkTextbox(description_frame, width=300, height=80)
        description_text.pack(fill=ctk.BOTH, expand=True)
        description_text.insert("1.0", document.get('description', ''))
        
        # Variables spécifiques du modèle
        variables_frame = ctk.CTkFrame(main_frame)
        variables_frame.pack(fill=ctk.BOTH, expand=True, pady=10)
        
        ctk.CTkLabel(variables_frame, text="Variables spécifiques", font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        # Cadre défilable pour les variables
        variables_scroll = ctk.CTkScrollableFrame(variables_frame)
        variables_scroll.pack(fill=ctk.BOTH, expand=True, padx=10, pady=10)
        
        # Dictionnaire pour stocker les widgets des variables
        variable_widgets = {}
        
        # Récupérer les variables existantes
        existing_variables = document.get('variables', {})
        
        # Créer les widgets pour chaque variable
        variables = template.get('variables', [])
        
        if variables:
            ctk.CTkLabel(
                variables_scroll, 
                text="Complétez les informations spécifiques à ce document:",
                font=ctk.CTkFont(weight="bold")
            ).pack(anchor="w", pady=(0, 10))
            
            for i, var_name in enumerate(variables):
                # Ignorer les variables standard qui sont remplies automatiquement
                if var_name in ['client_name', 'client_company', 'client_email', 'client_phone', 'client_address']:
                    continue
                
                var_frame = ctk.CTkFrame(variables_scroll)
                var_frame.pack(fill=ctk.X, pady=5)
                
                ctk.CTkLabel(var_frame, text=f"{var_name}:").pack(side=ctk.LEFT, padx=5)
                
                # Récupérer la valeur existante
                existing_value = existing_variables.get(var_name, '')
                
                if "montant" in var_name.lower() or "prix" in var_name.lower() or "cout" in var_name.lower() or "coût" in var_name.lower():
                    # Champ monétaire
                    money_frame = ctk.CTkFrame(var_frame, fg_color="transparent")
                    money_frame.pack(side=ctk.RIGHT, fill=ctk.X, expand=True)
                    
                    var = ctk.StringVar(value=existing_value)
                    entry = ctk.CTkEntry(money_frame, textvariable=var, width=150)
                    entry.pack(side=ctk.LEFT, fill=ctk.X, expand=True, padx=5)
                    
                    ctk.CTkLabel(money_frame, text="€").pack(side=ctk.LEFT)
                    
                    variable_widgets[var_name] = {"widget": entry, "var": var, "type": "money"}
                elif "date" in var_name.lower():
                    # Champ date
                    var = ctk.StringVar(value=existing_value)
                    entry = ctk.CTkEntry(var_frame, textvariable=var, width=150)
                    entry.pack(side=ctk.RIGHT, fill=ctk.X, expand=True, padx=5)
                    
                    variable_widgets[var_name] = {"widget": entry, "var": var, "type": "date"}
                else:
                    # Champ texte standard
                    var = ctk.StringVar(value=existing_value)
                    entry = ctk.CTkEntry(var_frame, textvariable=var, width=250)
                    entry.pack(side=ctk.RIGHT, fill=ctk.X, expand=True, padx=5)
                    
                    variable_widgets[var_name] = {"widget": entry, "var": var, "type": "text"}
        else:
            ctk.CTkLabel(
                variables_scroll, 
                text="Ce modèle ne contient pas de variables personnalisables.",
                text_color="gray"
            ).pack(anchor="w", pady=10)
        
        # Boutons
        button_frame = ctk.CTkFrame(dialog)
        button_frame.pack(fill=ctk.X, pady=10, padx=20)
        
        # Fonction pour mettre à jour le document
        def update_document_data():
            # Récupérer le titre et la date
            title = title_var.get().strip()
            if not title:
                DialogUtils.show_message(dialog, "Attention", "Veuillez entrer un titre", "warning")
                return
            
            date = date_var.get().strip()
            description = description_text.get("1.0", "end-1c").strip()
            
            # Récupérer les valeurs des variables
            variables_values = {}
            for var_name, widget_info in variable_widgets.items():
                if widget_info["type"] in ["text", "money", "date"]:
                    variables_values[var_name] = widget_info["var"].get().strip()
            
            # Mettre à jour l'objet document
            document['title'] = title
            document['date'] = date
            document['description'] = description
            document['variables'] = variables_values
            document['updated_at'] = datetime.now().isoformat()
            
            # Générer le contenu du document en remplaçant les variables
            content = self.generate_document_content(template, client, variables_values)
            
            # Vérifier si le fichier existe
            file_path = document.get('file_path')
            
            if not file_path or not os.path.exists(file_path):
                # Générer un nouveau fichier
                doc_filename = self.generate_filename(template.get("type", "document"), client.get('name', 'client'), date)
                format_type = self.model.config.get("document.default_format", "pdf")
                doc_filepath = os.path.join(self.model.paths['documents'], f"{doc_filename}.{format_type}")
                document['file_path'] = doc_filepath
            else:
                doc_filepath = file_path
            
            # Déterminer le format
            _, ext = os.path.splitext(doc_filepath)
            format_type = ext.lower().strip('.')
            
            # Générer le fichier selon le format
            try:
                if format_type == "pdf":
                    self.generate_pdf(doc_filepath, content, title, client, self.model.config.get("app.company_name", ""))
                else:
                    self.generate_docx(doc_filepath, content, title, client, self.model.config.get("app.company_name", ""))
                
                # Sauvegarder les documents
                self.model.save_documents()
                
                # Ajouter l'activité
                self.model.add_activity("document", f"Document mis à jour: {title}")
                
                # Mettre à jour la vue
                self.view.update_view()
                
                # Fermer la fenêtre
                dialog.destroy()
                
                # Afficher un message de succès
                DialogUtils.show_message(self.view.parent, "Succès", f"Document '{title}' mis à jour avec succès", "success")
                
                logger.info(f"Document mis à jour: {document_id} - {title}")
                
            except Exception as e:
                logger.error(f"Erreur lors de la mise à jour du document: {e}")
                DialogUtils.show_message(dialog, "Erreur", f"Erreur lors de la mise à jour du document: {str(e)}", "error")
        
        # Bouton Annuler
        cancel_btn = ctk.CTkButton(
            button_frame,
            text="Annuler",
            command=dialog.destroy,
            width=120
        )
        cancel_btn.pack(side=ctk.RIGHT, padx=10)
        
        # Bouton Mettre à jour
        update_btn = ctk.CTkButton(
            button_frame,
            text="Mettre à jour",
            command=update_document_data,
            width=120
        )
        update_btn.pack(side=ctk.RIGHT, padx=10)
        
        logger.info("Formulaire de mise à jour de document affiché")
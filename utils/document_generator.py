import os
import re
import logging
from datetime import datetime
import io

# Pour les documents PDF
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm, inch
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer, Image
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab n'est pas installé. La génération de PDF sera limitée.")

# Pour les documents Word
try:
    import docx
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx n'est pas installé. La génération de DOCX sera limitée.")

logger = logging.getLogger("VynalDocsAutomator.DocumentGenerator")

class DocumentGenerator:
    """
    Générateur de documents pour l'application
    Version améliorée avec gestion d'erreur robuste
    """
    
    def __init__(self, config_manager=None):
        """
        Initialise le générateur de documents
        
        Args:
            config_manager: Gestionnaire de configuration
        """
        self.config = config_manager
        logger.info("DocumentGenerator initialisé avec les capacités suivantes:")
        logger.info(f"- PDF (ReportLab): {'Disponible' if REPORTLAB_AVAILABLE else 'Non disponible'}")
        logger.info(f"- DOCX (python-docx): {'Disponible' if DOCX_AVAILABLE else 'Non disponible'}")
    
    def strip_html(self, html_content):
        """
        Enlève les balises HTML du contenu tout en préservant le texte
        
        Args:
            html_content: Contenu HTML
            
        Returns:
            str: Texte sans balises HTML
        """
        # Si le contenu n'est pas une chaîne, le convertir
        if not isinstance(html_content, str):
            return str(html_content)
        
        # Traiter certaines balises spéciales pour conserver la mise en forme
        # Remplacer les listes à puces par des caractères spéciaux
        html_content = re.sub(r'<li>(.*?)</li>', r'• \1\n', html_content)
        
        # Remplacer les sauts de ligne HTML par des sauts de ligne réels
        html_content = html_content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        
        # Remplacer les paragraphes par des lignes avec double saut de ligne
        html_content = re.sub(r'<p>(.*?)</p>', r'\1\n\n', html_content)
        
        # Remplacer les titres avec mise en forme
        html_content = re.sub(r'<h1>(.*?)</h1>', r'\1\n\n', html_content)
        html_content = re.sub(r'<h2>(.*?)</h2>', r'\1\n\n', html_content)
        html_content = re.sub(r'<h3>(.*?)</h3>', r'\1\n\n', html_content)
        
        # Remplacer les balises de style courantes
        html_content = re.sub(r'<strong>(.*?)</strong>', r'\1', html_content)
        html_content = re.sub(r'<b>(.*?)</b>', r'\1', html_content)
        html_content = re.sub(r'<em>(.*?)</em>', r'\1', html_content)
        html_content = re.sub(r'<i>(.*?)</i>', r'\1', html_content)
        html_content = re.sub(r'<u>(.*?)</u>', r'\1', html_content)
        
        # Supprimer les balises de listes
        html_content = html_content.replace('<ul>', '').replace('</ul>', '\n')
        html_content = html_content.replace('<ol>', '').replace('</ol>', '\n')
        
        # Supprimer toutes les autres balises HTML
        html_content = re.sub(r'<[^>]*>', '', html_content)
        
        # Supprimer les espaces multiples
        html_content = re.sub(r' +', ' ', html_content)
        
        # Supprimer les sauts de ligne multiples
        html_content = re.sub(r'\n{3,}', '\n\n', html_content)
        
        # Gérer les entités HTML courantes
        html_entities = {
            '&amp;': '&', 
            '&lt;': '<', 
            '&gt;': '>', 
            '&quot;': '"', 
            '&apos;': "'",
            '&nbsp;': ' '
        }
        
        for entity, char in html_entities.items():
            html_content = html_content.replace(entity, char)
        
        return html_content.strip()
    
    def replace_variables(self, content, variables):
        """
        Remplace les variables dans un contenu de manière sécurisée
        Supporte à la fois les formats {variable} et {{variable}}
        
        Args:
            content: Contenu avec variables
            variables: Dictionnaire des variables et leurs valeurs
            
        Returns:
            str: Contenu avec variables remplacées
        """
        import re
        
        # Si le contenu n'est pas une chaîne, le convertir
        if not isinstance(content, str):
            content = str(content)
        
        # Préparation des valeurs pour éviter les valeurs None
        safe_variables = {}
        for key, value in variables.items():
            if value is None:
                safe_variables[key] = ""
            else:
                safe_variables[key] = str(value)
        
        # Remplacer d'abord les variables au format {{variable}}
        for var_name, var_value in safe_variables.items():
            # Format {{variable}}
            pattern = r"{{" + re.escape(var_name) + r"}}"
            content = re.sub(pattern, var_value, content)
        
        # Ensuite remplacer les variables au format {variable}
        for var_name, var_value in safe_variables.items():
            # Format {variable} (pour la rétrocompatibilité)
            # Utiliser une expression régulière pour éviter les faux positifs
            pattern = r"{" + re.escape(var_name) + r"}"
            content = re.sub(pattern, var_value, content)
        
        # Rechercher les variables non remplacées au format {{variable}}
        remaining_vars = re.findall(r'{{([^{}]*?)}}', content)
        if remaining_vars:
            logger.warning(f"Variables non remplacées (format {{variable}}): {remaining_vars}")
            # Remplacer les variables non trouvées par une chaîne vide
            for var in remaining_vars:
                content = content.replace(f"{{{{{var}}}}}", "")
        
        # Rechercher les variables non remplacées au format {variable}
        remaining_simple_vars = re.findall(r'{([^{}]*)}', content)
        if remaining_simple_vars:
            logger.warning(f"Variables non remplacées (format {variable}): {remaining_simple_vars}")
            # Remplacer les variables non trouvées par une chaîne vide
            for var in remaining_simple_vars:
                content = content.replace(f"{{{var}}}", "")
        
        return content
    
    def clean_filename(self, name):
        """
        Nettoie un nom pour qu'il soit utilisable dans un nom de fichier
        
        Args:
            name: Nom à nettoyer
            
        Returns:
            str: Nom nettoyé
        """
        # Supprimer les caractères spéciaux et remplacer les espaces par des underscores
        name = re.sub(r'[\\/*?:"<>|]', '', name)
        name = name.replace(' ', '_')
        name = name.replace('/', '_')
        name = name.replace('\\', '_')
        name = name.strip('.')  # Supprime les points en début et fin
        
        return name
    
    def generate_document(self, file_path, template, client, company_info, variables, format_type=None):
        """
        Génère un document à partir d'un modèle avec gestion d'erreur robuste
        
        Args:
            file_path: Chemin du fichier à créer
            template: Modèle de document (dict)
            client: Informations du client (dict)
            company_info: Informations de l'entreprise (dict)
            variables: Variables spécifiques pour le document
            format_type: Format du document (pdf, docx ou txt)
            
        Returns:
            bool: True si le document a été généré avec succès, False sinon
        """
        try:
            # Déterminer le format si non spécifié
            if format_type is None:
                format_type = "pdf"
                if self.config:
                    format_type = self.config.get("document.default_format", "pdf")
            
            format_type = format_type.lower()
            
            # Vérifier la disponibilité des modules nécessaires
            if format_type == "pdf" and not REPORTLAB_AVAILABLE:
                logger.warning("ReportLab n'est pas installé, utilisation du format TXT à la place")
                format_type = "txt"
            elif format_type == "docx" and not DOCX_AVAILABLE:
                logger.warning("python-docx n'est pas installé, utilisation du format TXT à la place")
                format_type = "txt"
            
            # S'assurer que l'extension correspond au format
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext[1:] != format_type:
                file_path = os.path.splitext(file_path)[0] + f".{format_type}"
            
            # S'assurer que le répertoire parent existe
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Préparer les variables
            all_variables = {}
            
            # Variables standard du client
            all_variables.update({
                "client_name": client.get("name", ""),
                "client_company": client.get("company", ""),
                "client_email": client.get("email", ""),
                "client_phone": client.get("phone", ""),
                "client_address": client.get("address", "")
            })
            
            # Variables de l'entreprise
            all_variables.update({
                "company_name": company_info.get("name", ""),
                "company_address": company_info.get("address", ""),
                "company_email": company_info.get("email", ""),
                "company_phone": company_info.get("phone", ""),
                "company_website": company_info.get("website", "")
            })
            
            # Variables spécifiques
            all_variables.update(variables)
            
            # Date actuelle
            all_variables["date"] = datetime.now().strftime("%Y-%m-%d")
            
            # Obtenir le contenu du modèle
            content = template.get("content", "")
            
            # Si le modèle a un chemin de fichier, l'utiliser
            if "file_path" in template and template["file_path"]:
                try:
                    template_file = template["file_path"]
                    if not os.path.isabs(template_file):
                        # Construire le chemin absolu si le chemin est relatif
                        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                        template_file = os.path.join(base_dir, "data", "templates", template_file)
                    
                    if os.path.exists(template_file):
                        with open(template_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                except Exception as e:
                    logger.error(f"Erreur lors de la lecture du fichier modèle: {e}")
                    # Continuer avec le contenu stocké dans la base de données
            
            # Remplacer les variables dans le contenu
            content = self.replace_variables(content, all_variables)
            
            # Titre du document
            title = template.get("name", "Document")
            if "title" in variables:
                title = variables["title"]
            
            # Obtenir le chemin du logo
            logo_path = None
            if self.config:
                logo_path = self.config.get("app.company_logo", "")
                
                # Vérifier que le logo existe
                if logo_path and not os.path.exists(logo_path):
                    logger.warning(f"Logo spécifié mais introuvable: {logo_path}")
                    logo_path = None
            
            # Générer le document selon le format avec gestion d'erreur
            try:
                if format_type == "pdf":
                    success = self.generate_pdf(file_path, content, title, client, company_info, logo_path)
                elif format_type == "docx":
                    success = self.generate_docx(file_path, content, title, client, company_info, logo_path)
                else:
                    success = self.generate_txt(file_path, content, title, client, company_info)
                
                if not success:
                    raise Exception(f"Échec lors de la génération au format {format_type}")
                
                logger.info(f"Document généré avec succès: {file_path}")
                return True
            
            except Exception as format_error:
                logger.error(f"Erreur lors de la génération au format {format_type}: {format_error}")
                
                # Solution de repli: générer un document texte
                try:
                    txt_path = os.path.splitext(file_path)[0] + ".txt"
                    success = self.generate_txt(txt_path, content, title, client, company_info)
                    
                    if success:
                        logger.info(f"Document texte de secours généré: {txt_path}")
                        return True
                    else:
                        raise Exception("Échec de la génération du document texte de secours")
                
                except Exception as fallback_error:
                    logger.error(f"Erreur lors de la génération du document texte de secours: {fallback_error}")
                    
                    # Dernière tentative: créer un document texte minimal
                    try:
                        minimal_path = os.path.splitext(file_path)[0] + ".txt"
                        with open(minimal_path, 'w', encoding='utf-8') as f:
                            f.write(f"Titre: {title}\n")
                            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
                            f.write(f"Client: {client.get('name', '')}\n\n")
                            f.write("Une erreur est survenue lors de la génération du document.\n")
                            f.write("Ce document est une version de secours minimale.\n")
                        
                        logger.info(f"Document texte minimal créé: {minimal_path}")
                        return True
                    
                    except Exception as minimal_error:
                        logger.error(f"Erreur lors de la création du document minimal: {minimal_error}")
                        return False
        
        except Exception as e:
            logger.error(f"Erreur globale lors de la génération du document: {e}")
            
            try:
                # Tentative ultime: créer un fichier texte d'erreur
                error_path = os.path.join(os.path.dirname(file_path), f"error_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt")
                with open(error_path, 'w', encoding='utf-8') as f:
                    f.write(f"ERREUR DE GÉNÉRATION DE DOCUMENT\n")
                    f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Erreur: {str(e)}\n")
                
                logger.info(f"Fichier d'erreur créé: {error_path}")
                return False
            except:
                return False
    
    def generate_txt(self, file_path, content, title, client, company_info):
        """
        Génère un document texte simple en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            
        Returns:
            bool: True si réussi, False sinon
        """
        try:
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML
            clean_content = self.strip_html(content)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # En-tête
                f.write(f"Titre: {title}\n")
                f.write(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
                f.write(f"Généré par: {company_info.get('name', 'Vynal Docs Automator')}\n\n")
                
                # Informations du client
                f.write("=== INFORMATIONS CLIENT ===\n")
                f.write(f"Nom: {client.get('name', '')}\n")
                if client.get('company'):
                    f.write(f"Entreprise: {client.get('company')}\n")
                f.write(f"Email: {client.get('email', '')}\n")
                if client.get('phone'):
                    f.write(f"Téléphone: {client.get('phone')}\n")
                if client.get('address'):
                    f.write(f"Adresse: {client.get('address')}\n")
                f.write("\n")
                
                # Contenu
                f.write("=== CONTENU ===\n")
                f.write(clean_content)
                
                # Pied de page
                f.write("\n\n")
                f.write("Document généré le " + datetime.now().strftime("%Y-%m-%d à %H:%M:%S"))
                f.write("\n")
                f.write(f"© {datetime.now().year} {company_info.get('name', 'Vynal Docs Automator')}")
            
            return True
        
        except Exception as e:
            logger.error(f"Erreur lors de la génération du fichier texte: {e}")
            return False
    
    def generate_pdf(self, file_path, content, title, client, company_info, logo_path=None):
        """
        Génère un document PDF en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            logo_path: Chemin du logo de l'entreprise
            
        Returns:
            bool: True si réussi, False sinon
        """
        try:
            if not REPORTLAB_AVAILABLE:
                logger.error("ReportLab n'est pas installé, impossible de générer un PDF")
                return False
            
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML
            clean_content = self.strip_html(content)
            
            # Créer un PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4
            
            # En-tête avec logo
            y_position = height - 50
            
            # Ajouter le logo si existant
            if logo_path and os.path.exists(logo_path):
                try:
                    logo_width = 150
                    logo_height = 50
                    c.drawImage(logo_path, 50, height - 70, width=logo_width, height=logo_height, preserveAspectRatio=True)
                    y_position = height - 100
                except Exception as logo_error:
                    logger.warning(f"Erreur lors de l'ajout du logo: {logo_error}")
            
            # Nom de l'entreprise
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y_position, company_info.get("name", ""))
            y_position -= 20
            
            # Informations générales
            c.setFont("Helvetica", 10)
            c.drawString(50, y_position, "Document généré par Vynal Docs Automator")
            y_position -= 15
            c.drawString(50, y_position, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            y_position -= 15
            
            # Ligne séparatrice
            c.line(50, y_position, width - 50, y_position)
            y_position -= 20
            
            # Titre du document
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y_position, title)
            y_position -= 30
            
            # Informations du client
            c.setFont("Helvetica", 12)
            c.drawString(50, y_position, f"Client: {client.get('name', '')}")
            y_position -= 15
            
            if client.get('company'):
                c.drawString(50, y_position, f"Entreprise: {client.get('company')}")
                y_position -= 15
            
            c.drawString(50, y_position, f"Email: {client.get('email', '')}")
            y_position -= 15
            
            if client.get('phone'):
                c.drawString(50, y_position, f"Téléphone: {client.get('phone')}")
                y_position -= 15
            
            if client.get('address'):
                c.drawString(50, y_position, f"Adresse: {client.get('address')}")
                y_position -= 15
            
            # Ligne séparatrice
            c.line(50, y_position, width - 50, y_position)
            y_position -= 20
            
            # Contenu du document
            # Diviser le contenu en lignes
            lines = clean_content.split('\n')
            c.setFont("Helvetica", 10)
            
            for line in lines:
                # Vérifier s'il reste assez d'espace sur la page
                if y_position < 50:  # Marge de bas de page de 50 points
                    c.showPage()  # Nouvelle page
                    c.setFont("Helvetica", 10)
                    y_position = height - 50
                
                # Dessiner la ligne
                c.drawString(50, y_position, line)
                y_position -= 15  # Espacement des lignes
            
            # Pied de page
            c.setFont("Helvetica", 8)
            c.drawString(width/2 - 100, 20, f"© {datetime.now().year} {company_info.get('name', '')}")
            
            # Finaliser le document
            c.showPage()
            c.save()
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du PDF: {e}")
            return False
    
    def generate_docx(self, file_path, content, title, client, company_info, logo_path=None):
        """
        Génère un document DOCX en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier à créer
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            logo_path: Chemin du logo de l'entreprise
            
        Returns:
            bool: True si réussi, False sinon
        """
        try:
            if not DOCX_AVAILABLE:
                logger.error("python-docx n'est pas installé, impossible de générer un DOCX")
                return False
            
            # Créer le répertoire parent si nécessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML
            clean_content = self.strip_html(content)
            
            # Créer un document Word
            doc = docx.Document()
            
            # En-tête avec logo si disponible
            if logo_path and os.path.exists(logo_path):
                try:
                    # Ajout du logo dans l'en-tête
                    header = doc.sections[0].header
                    header_para = header.paragraphs[0]
                    run = header_para.add_run()
                    run.add_picture(logo_path, width=Cm(5))
                except Exception as e:
                    logger.warning(f"Erreur lors de l'ajout du logo: {e}")
            
            # Titre du document
            title_para = doc.add_heading(title, level=1)
            
            # Informations du client
            client_para = doc.add_paragraph()
            client_para.add_run("Client: ").bold = True
            client_para.add_run(client.get('name', ''))
            
            if client.get('company'):
                company_para = doc.add_paragraph()
                company_para.add_run("Entreprise: ").bold = True
                company_para.add_run(client.get('company'))
            
            email_para = doc.add_paragraph()
            email_para.add_run("Email: ").bold = True
            email_para.add_run(client.get('email', ''))
            
            if client.get('phone'):
                phone_para = doc.add_paragraph()
                phone_para.add_run("Téléphone: ").bold = True
                phone_para.add_run(client.get('phone'))
            
            if client.get('address'):
                address_para = doc.add_paragraph()
                address_para.add_run("Adresse: ").bold = True
                address_para.add_run(client.get('address'))
            
            # Ligne séparatrice
            doc.add_paragraph("_______________________________________________________________")
            
            # Contenu
            # Diviser le contenu en paragraphes
            paragraphs = clean_content.split('\n\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text)
            
            # Pied de page
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_run = footer_para.add_run(f"© {datetime.now().year} {company_info.get('name', '')} - Confidentiel")
            footer_run.font.size = Pt(8)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Enregistrer le document
            doc.save(file_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération du DOCX: {e}")
            return False